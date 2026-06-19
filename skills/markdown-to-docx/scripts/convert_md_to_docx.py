#!/usr/bin/env python3
"""
Professional Markdown to Word Document Converter
Direct MD → DOCX using python-docx (no PDF intermediate)
Produces clean, professional Word documents with proper formatting
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class MarkdownToDocx:
    def __init__(self, md_file, docx_file):
        self.md_file = md_file
        self.docx_file = docx_file
        self.doc = Document()
        self.set_default_styles()

    def set_default_styles(self):
        """Configure professional document styles"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(51, 51, 51)

    def clean_markdown(self, text):
        """Remove problematic characters - KEEP emoji icons"""
        # Remove only truly problematic control characters and null bytes
        # Keep emoji icons - they render as visual icons in Word
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
        return text

    def parse_markdown(self, md_text):
        """Parse markdown into blocks with type and content"""
        blocks = []
        lines = md_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Skip empty lines and separators (---)
            if not line.strip() or line.strip() == '---':
                i += 1
                continue

            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                content = line.lstrip('#').strip()
                blocks.append({'type': 'heading', 'level': level, 'content': content})
                i += 1

            # Tables
            elif line.strip().startswith('|'):
                table_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                blocks.append({'type': 'table', 'lines': table_lines})

            # Bullet/numbered lists
            elif re.match(r'^[\s]*[-*][\s]', line):
                list_items = []
                while i < len(lines) and (re.match(r'^[\s]*[-*][\s]', lines[i]) or lines[i].startswith('  ')):
                    list_items.append(lines[i])
                    i += 1
                blocks.append({'type': 'bullet_list', 'items': list_items})

            # Code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1  # Skip closing ```
                blocks.append({'type': 'code', 'lines': code_lines})

            # Paragraphs
            else:
                para_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not re.match(r'^[\s]*[-*]', lines[i]) and not lines[i].startswith('|') and not lines[i].startswith('```') and lines[i].strip() != '---':
                    para_lines.append(lines[i])
                    i += 1
                blocks.append({'type': 'paragraph', 'lines': para_lines})

        return blocks

    def format_inline(self, text):
        """Apply inline formatting (bold, italic, links)"""
        # Apply formatting FIRST
        # Replace markdown bold
        text = re.sub(r'\*\*(.*?)\*\*', r'<BOLD>\1</BOLD>', text)
        text = re.sub(r'__(.*?)__', r'<BOLD>\1</BOLD>', text)

        # Replace markdown italic
        text = re.sub(r'(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)', r'<ITALIC>\1</ITALIC>', text)

        # Replace links
        text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<LINK>\1||\2</LINK>', text)

        # Replace emoji with simple Unicode symbols that render everywhere
        text = text.replace('✅', '✓')
        text = text.replace('❌', '✗')
        text = text.replace('📁', '▶')
        text = text.replace('⚠️', '⚠')
        text = text.replace('🔄', '↻')
        text = text.replace('🔴', '●')
        text = text.replace('🟢', '●')
        text = text.replace('⏳', '⧗')
        text = text.replace('❓', '?')

        return text

    def _add_formatted_text(self, paragraph, text, size=11):
        """Helper: add formatted text (bold, italic, links) to a paragraph"""
        pattern = r'(<BOLD>.*?</BOLD>|<ITALIC>.*?</ITALIC>|<LINK>.*?</LINK>|[^<]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            part = ''.join(char if ord(char) >= 32 or char in '\n\r\t' else '' for char in part)

            if part.startswith('<BOLD>') and part.endswith('</BOLD>'):
                content = part[6:-7]
                content = ''.join(char if ord(char) >= 32 or char in '\n\r\t' else '' for char in content)
                if content.strip():
                    run = paragraph.add_run(content)
                    run.bold = True
                    run.font.size = Pt(size)

            elif part.startswith('<ITALIC>') and part.endswith('</ITALIC>'):
                content = part[8:-9]
                content = ''.join(char if ord(char) >= 32 or char in '\n\r\t' else '' for char in content)
                if content.strip():
                    run = paragraph.add_run(content)
                    run.italic = True
                    run.font.size = Pt(size)

            elif part.startswith('<LINK>') and part.endswith('</LINK>'):
                link_content = part[6:-7]
                if '||' in link_content:
                    text_part, url = link_content.split('||', 1)
                    text_part = text_part.strip()
                    url = url.strip()
                    if text_part and url:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        run = paragraph.add_run(text_part)
                        run.font.size = Pt(size)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                        rId = paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                        r = run._element
                        hyperlink_xml = f'<w:hyperlink {nsdecls("w", "r")} r:id="{rId}"/>'
                        hyperlink = parse_xml(hyperlink_xml)
                        r.getparent().replace(r, hyperlink)
                        hyperlink.append(r)

            elif part and part.strip():
                run = paragraph.add_run(part)
                run.font.size = Pt(size)

    def add_formatted_paragraph(self, text, style='Normal', size=11):
        """Add paragraph with inline formatting"""
        text = self.format_inline(text)
        p = self.doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(6)
        self._add_formatted_text(p, text, size=size)
        return p

    def add_heading(self, text, level):
        """Add heading with appropriate level"""
        text = self.clean_markdown(text)
        heading = self.doc.add_heading(text, level=level)
        # Keep headings black, not colored
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    def add_table(self, lines):
        """Parse and add table to document with formatting"""
        if len(lines) < 3:
            return

        # Parse table
        header_row = [cell.strip() for cell in lines[0].split('|')[1:-1]]
        rows = []
        for line in lines[2:]:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(row) == len(header_row):
                rows.append(row)

        if not rows:
            return

        # Create table
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(header_row))
        table.style = 'Table Grid'

        # Header row with formatting
        for i, cell_text in enumerate(header_row):
            cell = table.rows[0].cells[i]
            cell_text = self.format_inline(cell_text)

            # Clear default paragraph
            cell.paragraphs[0].text = ''
            para = cell.paragraphs[0]

            # Add formatted content
            self._add_formatted_text(para, cell_text, size=10)

            # Format header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

            # Header background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E8EEF7')
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Data rows with formatting
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell_text = self.format_inline(cell_text)

                # Clear default paragraph
                cell.paragraphs[0].text = ''
                para = cell.paragraphs[0]

                # Add formatted content
                self._add_formatted_text(para, cell_text, size=10)

        self.doc.add_paragraph()  # Spacing after table

    def add_bullet_list(self, items):
        """Add bullet list with formatting"""
        for item in items:
            # Get indentation and bullet carefully
            stripped = item.lstrip(' \t')
            indent = len(item) - len(stripped)

            # Remove ONLY the bullet marker, preserve markdown
            if stripped.startswith('- '):
                content = stripped[2:]  # Remove "- "
            elif stripped.startswith('* '):
                content = stripped[2:]  # Remove "* "
            else:
                content = stripped

            content = content.strip()
            content = self.format_inline(content)

            level = min(indent // 2, 2)
            p = self.doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.5 + (level * 0.25))
            p.paragraph_format.first_line_indent = Inches(-0.25)

            self._add_formatted_text(p, content, size=11)

    def add_code_block(self, lines):
        """Add code block"""
        code_text = '\n'.join(lines)
        p = self.doc.add_paragraph(code_text, style='Normal')
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

    def convert(self):
        """Convert markdown file to Word document"""
        # Read markdown
        with open(self.md_file, 'r', encoding='utf-8') as f:
            md_text = f.read()

        # Parse into blocks
        blocks = self.parse_markdown(md_text)

        # Process each block
        for block in blocks:
            if block['type'] == 'heading':
                self.add_heading(block['content'], block['level'])

            elif block['type'] == 'table':
                self.add_table(block['lines'])

            elif block['type'] == 'bullet_list':
                self.add_bullet_list(block['items'])

            elif block['type'] == 'code':
                self.add_code_block(block['lines'])

            elif block['type'] == 'paragraph':
                text = '\n'.join(block['lines'])
                self.add_formatted_paragraph(text)

        # Save document
        self.doc.save(self.docx_file)
        print(f"✓ Document created: {self.docx_file}")
        print(f"  Size: {Path(self.docx_file).stat().st_size:,} bytes")

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 convert_md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2]

    converter = MarkdownToDocx(md_file, docx_file)
    converter.convert()

if __name__ == '__main__':
    main()
